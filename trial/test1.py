import sqlite3
from docx import Document

def fetch_questions(class_name, subject, assessment, chapters, standards):
    """
    Fetch questions from the database based on the provided criteria.
    """
    conn = sqlite3.connect('main.db')
    cursor = conn.cursor()
    placeholders = ', '.join(['?'] * len(chapters))  # For multiple chapters
    standards_placeholders = ', '.join(['?'] * len(standards))
    query = f'''
        SELECT id, question_text, marks, chapter, standard 
        FROM questions 
        WHERE class = ? AND subject = ? AND assessment = ? 
        AND chapter IN ({placeholders}) 
        AND standard IN ({standards_placeholders})
    '''
    cursor.execute(query, (class_name, subject, assessment, *chapters, *standards))
    questions = cursor.fetchall()
    conn.close()
    return questions


def distribute_questions_by_standard(questions, standard_distribution, total_questions):
    """
    Distribute questions based on standard preferences or randomly if not specified.
    """
    grouped_questions = {}
    for question in questions:
        standard = question[4]
        if standard not in grouped_questions:
            grouped_questions[standard] = []
        grouped_questions[standard].append(question)

    selected_questions = []
    for standard, count in standard_distribution.items():
        if standard in grouped_questions:
            selected_questions.extend(grouped_questions[standard][:count])

    remaining_questions = [
        q for q in questions if q not in selected_questions
    ]
    if len(selected_questions) < total_questions:
        selected_questions.extend(remaining_questions[:total_questions - len(selected_questions)])
    
    return selected_questions


def save_question_paper_to_word(sections, output_file):
    """
    Save the question paper to a Word document.
    """
    document = Document()
    document.add_heading("Question Paper", level=1)

    for section_num, section in enumerate(sections, start=1):
        document.add_heading(f"Section {section_num}: {section['marks']} Marks", level=2)
        for question_num, question in enumerate(section['questions'], start=1):
            question_text = question[1]
            chapter = question[3]
            standard = question[4]
            document.add_paragraph(f"{question_num}. {question_text} (Chapter: {chapter}, Standard: {standard})")

    document.save(output_file)


def main():
    # Step 1: Gather user inputs
    class_name = input("Enter the class: ")
    subject = input("Enter the subject: ")
    assessment = input("Enter the assessment type: ")

    print("Enter chapter names (one per line). Type 'done' when finished:")
    chapters = []
    while True:
        chapter = input()
        if chapter.lower() == 'done':
            break
        chapters.append(chapter)

    num_sections = int(input("Enter the number of sections: "))
    sections = []

    for i in range(num_sections):
        print(f"\nSection {i+1}:")
        marks = int(input("Enter marks for this section: "))
        num_questions = int(input("Enter the number of questions for this section: "))

        print("Enter the standards for this section (e.g., L1, L2). Type 'done' when finished:")
        standards = []
        while True:
            standard = input()
            if standard.lower() == 'done':
                break
            standards.append(standard)
        print("st",standard)

        print("Specify number of questions per standard (e.g., 3 for L1, 2 for L2). Type 'done' to finish:")
        standard_distribution = {}
        while True:
            entry = input("Standard and count (e.g., L1 3): ")
            if entry.lower() == 'done':
                break
            standard, count = entry.split()
            standard_distribution[standard] = int(count)
        
        print("stqu",standard_distribution)

        # Fetch questions for this section
        questions = fetch_questions(class_name, subject, assessment, chapters, standards)

        if len(questions) < num_questions:
            print(f"Only {len(questions)} questions available for this section. Adjusting to available questions.")
            num_questions = len(questions)

        selected_questions = distribute_questions_by_standard(questions, standard_distribution, num_questions)

        sections.append({
            'marks': marks,
            'questions': selected_questions
        })

    # Step 2: Print the question paper
    print("\n=======================")
    print("       QUESTION PAPER       ")
    print("=======================\n")
    for section_num, section in enumerate(sections, start=1):
        print(f"\nSection {section_num}: {section['marks']} Marks")
        print("-" * 30)
        for question_num, question in enumerate(section['questions'], start=1):
            question_text = question[1]
            chapter = question[3]
            standard = question[4]
            print(f"{question_num}. {question_text} (Chapter: {chapter}, Standard: {standard})")
        print("\n")

    # Step 3: Save to Word file
    output_file = "Question_Paper.docx"
    save_question_paper_to_word(sections, output_file)
    print(f"Question paper saved to {output_file}")


if __name__ == "__main__":
    main()
