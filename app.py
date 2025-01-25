from flask import Flask, request, render_template, jsonify , send_from_directory , send_file,session,flash,url_for,redirect
import sqlite3
import os
from werkzeug.utils import secure_filename
import google.generativeai as genai
import pandas as pd
from db import init_main_database , save_file_to_main_db , get_all_files_from_main_db, get_all_questions_based_on_fileid_from_main_db,update_question_in_db,check_question_exists_in_main_db,insert_single_question_to_main_db,insert_question_to_main_db,get_file_id_by_filename,delete_file_by_id ,get_all_questions_based_on_paperid_from_main_db
from db import get_only_questions_based_on_fileid_from_main_db ,insert_chapters_to_main_db ,insert_academic_standard_to_main_db,get_dropdown_data,get_chapters , fetch_questions,distribute_questions_by_standard , insert_generatedpaper_to_main_db , insert_generatedquestion_to_main_db,get_all_papers_from_main_db ,replace_question_in_db,delete_generatedpaper_from_db,delete_questionpaper_from_db
from db import get_teacherfeedback_from_main_db , insert_teacherfeedback,get_user_by_username,add_user,get_all_questions_based_on_userid_from_main_db,get_question_based_on_id_from_main_db ,delete_question_from_db,get_all_files_based_on_userid_from_main_db,get_all_questions_with_all_userids_from_main_db,get_all_papers_based_on_userid_from_main_db
from db import get_all_users , add_admin_user,edit_user_based_on_id ,delete_user_based_on_id,get_user_by_id , get_files_uploaded_count , get_questionpapers_count
import json
import re
from datetime import datetime
from docx import Document
import pandas as pd
import numpy as np
from sklearn.preprocessing import MinMaxScaler
import matplotlib.pyplot as plt
import io
import base64
import random

# from db1 import get_dropdown , get_questions_from_db , rank_questions_by_sentence_transformer , rank_questions_with_bert




# Configure Gemini API
GOOGLE_API_KEY = 'AIzaSyDvpbV34TFp2hGXWF5Hl9zONjlKeKoAuv8'
genai.configure(api_key=GOOGLE_API_KEY)
# model = genai.GenerativeModel('gemini-pro')

# Initialize the model (Gemini)
model = genai.GenerativeModel('gemini-1.5-flash')

app = Flask(__name__)

# Configure upload folder and allowed extensions
UPLOAD_FOLDER = 'uploads/'
ALLOWED_EXTENSIONS = {'csv', 'xlsx', 'docx', 'pdf', 'png', 'jpeg', 'jpg'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
DIAGRAMS_FOLDER = 'diagrams/'
app.config['DIAGRAMS_FOLDER'] = DIAGRAMS_FOLDER
PAPERS_FOLDER = 'papers/'
app.config['PAPERS_FOLDER'] = PAPERS_FOLDER

EXCELUPLOAD_FOLDER = 'excel_upload/'
ALLOWED_EXTENSIONS1 = {'csv'}

app.config['EXCELUPLOAD_FOLDER'] = EXCELUPLOAD_FOLDER
app.secret_key = 'your_secret_key'

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS1


# Initialize databases
init_main_database()

add_admin_user()

# Function to check if file extension is allowed
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS



# Redirect to login page if not logged in
@app.before_request
def require_login():
    # Exclude the login route from this check
    if 'user_id' not in session and request.endpoint not in ['login', 'static']:
        return redirect(url_for('login'))
    


# Login route
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        user = get_user_by_username(username)  # Use the function from db.py
        print(user)
        if user and user[1] == password:  # Direct password comparison
            session['user_id'] = user[0]
            session['username'] = user[1]
            session['role'] = user[3]
            return redirect(url_for('index'))  # Redirect to home page
        else:
            flash('Invalid username or password')
    return render_template('login.html')


# Logout route
@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

# Role-based route restriction
def role_required(allowed_roles):
    def decorator(f):
        def wrapper(*args, **kwargs):
            if 'role' not in session or session['role'] not in allowed_roles:
                flash('Access denied.')
                return redirect(url_for('login'))
            return f(*args, **kwargs)
        wrapper.__name__ = f.__name__
        return wrapper
    return decorator



# @app.route("/rankedquestions", methods=["GET", "POST"])
# def RankedQuestions():
#     classes, subjects = get_dropdown()
#     ranked_questions = []
#     query_ranked_questions = []

#     if request.method == "POST":
#         selected_class = request.form.get("class")
#         selected_subject = request.form.get("subject")
#         query_text = request.form.get("query_text")

#         questions, available_columns = get_questions_from_db()
#         filtered_questions = [q for q in questions if q[available_columns.index("class")] == selected_class and q[available_columns.index("subject")] == selected_subject]

#         if filtered_questions:
#             ranked_questions = rank_questions_by_sentence_transformer(filtered_questions, available_columns)

#         if query_text:
#             query_ranked_questions = rank_questions_with_bert(query_text)

#     return render_template("RankedQuestions.html", classes=classes, subjects=subjects, ranked_questions=ranked_questions, query_ranked_questions=query_ranked_questions)




@app.route('/account')
def account():
    userid = session.get('user_id') 
    role = session.get('role')  
 
    try:
        user = get_user_by_id(userid)
        uploaded_files_count = get_files_uploaded_count(userid)
        questionpapers_count = get_questionpapers_count(userid)

        if user:
            user_data = {
                'id': user[0],
                'username': user[1],
                'password': user[2],  # Again, don't send passwords in plaintext
                'role': user[3],
                'designation': user[4],
                'qualification': user[5],
                'class_teacher': user[6],
                'uploaded_files_count': uploaded_files_count,
                'generated_papers_count': questionpapers_count
            }
            # Pass the user_data to the template instead of JSON response
            return render_template('about.html', user_details=user_data,role=role)
        else:
            return jsonify({'error': 'User not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 400

# Example route for adding a user
@app.route('/add_user', methods=['POST'])
@role_required(['admin'])
def add_user_route():
    username = request.form['username']
    password = request.form['password']
    role = request.form['role']
    designation = request.form['designation']
    qualification = request.form['qualification']
    class_teacher = request.form['classTeacher']
    try:
        add_user(username, password, role, designation, qualification, class_teacher)  # Add user to the database
        flash('User added successfully')
        return jsonify({'success': True, 'message': 'User added successfully'}), 200  # Success
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 400  # Error


@app.route('/addTeachers')
@role_required(['admin'])
def add_teachers():
    role = session.get('role')  
    if not role:  # Redirect to login if not authenticated
        return redirect(url_for('login'))
    return render_template('AddUsers.html',role=role)

@app.route('/get_teachers', methods=['GET'])
@role_required(['admin'])
def get_users():
    try:
        users = get_all_users()
        # Convert user data into a dictionary and return as JSON
        users_list = [{
            'id': user[0],
            'username': user[1],
            'password': user[2],  # Note: In production, don't send passwords in plaintext!
            'role': user[3],
            'designation': user[4],
            'qualification': user[5],
            'class_teacher': user[6]
        } for user in users]
        
        return jsonify(users_list), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 400


@app.route('/get_teachers/<int:user_id>', methods=['GET'])
@role_required(['admin'])
def get_user(user_id):
    try:
        user = get_user_by_id(user_id)
        uploaded_files_count = get_files_uploaded_count(user_id)
        questionpapers_count = get_questionpapers_count(user_id)
        if user:
            user_data = {
                'id': user[0],
                'username': user[1],
                'password': user[2],  # Again, don't send passwords in plaintext
                'role': user[3],
                'designation': user[4],
                'qualification': user[5],
                'class_teacher': user[6],
                'uploaded_files_count': uploaded_files_count,
                'generated_papers_count': questionpapers_count
            }
            return jsonify(user_data), 200
        else:
            return jsonify({'error': 'User not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 400


@app.route('/edit_user', methods=['POST'])
@role_required(['admin'])
def edit_user():
    user_id = request.form['userId']
    username = request.form['username']
    password = request.form['password']
    role = request.form['role']
    designation = request.form['designation']
    qualification = request.form['qualification']
    class_teacher = request.form['classTeacher']

    try:
        edit_user_based_on_id(username, password, role, designation, qualification, class_teacher, user_id)

        return jsonify({'success': True, 'message': 'User updated successfully'}), 200
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 400


@app.route('/delete_user', methods=['POST'])
@role_required(['admin'])
def delete_user():
    user_id = request.form['userId']
    
    try:
        delete_user_based_on_id(user_id)
        return jsonify({'success': True, 'message': 'User deleted successfully'}), 200
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 400


@app.route('/')
def index():
    role = session.get('role')  
    if not role:  # Redirect to login if not authenticated
        return redirect(url_for('login'))
    return render_template('index.html',role=role)


@app.route('/AddQuestions')
@role_required(['admin', 'teacher'])
def updateq():
    role = session.get('role')  
    if not role:  # Redirect to login if not authenticated
        return redirect(url_for('login'))
    return render_template('AddQuestions.html',role=role)

@app.route('/EditQuestions')
@role_required(['admin', 'teacher'])
def editq():
    role = session.get('role') 

    if not role:  # Redirect to login if not authenticated
        return redirect(url_for('login'))
    return render_template('EditQuestions1.html',role=role)

@app.route('/GenerateQuestion')
@role_required(['admin', 'teacher'])
def generatequestionpaper():
    role = session.get('role')  
    if not role:  # Redirect to login if not authenticated
        return redirect(url_for('login'))
    classes, subjects, assessments = get_dropdown_data()
    return render_template('GenerateQuestion.html', classes=classes, subjects=subjects, assessments=assessments,role=role)


@app.route('/GeneratedPapers')
@role_required(['admin', 'teacher'])
def loadquestiontemplate():
    role = session.get('role')  
    if not role:  # Redirect to login if not authenticated
        return redirect(url_for('login'))
    return render_template('Generatedpapers.html',role=role)



@app.route('/GeneratedPapers1')
@role_required(['admin', 'teacher'])
def loadquestionpapers():
    user_id = session.get('user_id') 
    role=session.get('role')
    if not user_id:  # Redirect to login if not authenticated
        return redirect(url_for('login'))

    if role=="admin":
        data = get_all_papers_from_main_db()
    else:
        data = get_all_papers_based_on_userid_from_main_db(user_id)  # Get questions based on user_id


    generated_papers = []
    for file in data:
        generated_papers.append({
            'id': file[0],
            'class': file[1],  
            'subject': file[2],
            'assessment': file[3],  
            'generated_time': file[5]
        })
    return jsonify(generated_papers)


@app.route('/get_generated_questions/<paper_id>', methods=['GET'])
@role_required(['admin', 'teacher'])
def get_generated_questions(paper_id):
    try:
        questions = get_all_questions_based_on_paperid_from_main_db(paper_id)  # Get questions based on file_id
       
     # Dictionary to group questions by section ID
        grouped_sections = {}

        for question in questions:
            section_id = question[3]
            if section_id not in grouped_sections:
                grouped_sections[section_id] = {
                    'sectionid': section_id,
                    'questions': []
                }
            
            grouped_sections[section_id]['questions'].append({
                'id': question[0],
                'question_text': question[1],
                'chapter': question[2],
                'marks': question[4],
                'standard': question[6],
                'diagram': question[5]
            })
        
        # Convert the grouped data into a list
        sections_data = list(grouped_sections.values())
        # print(sections_data)
        
        return jsonify(sections_data)  # Return grouped sections as JSON
    except Exception as e:
        return jsonify({'error': str(e)}), 500  # Handle unexpected errors
        


@app.route('/generate_new_question_with_ai', methods=['POST'])
@role_required(['admin', 'teacher'])
def edit_question_using_ai():
    try:
        # Access the JSON data sent from the frontend
        data = request.get_json()
        
        # Extract individual fields from the data
        class_label = data.get('class')
        subject_label = data.get('subject')
        marks_label = data.get('marks')
        standard_label = data.get('standard')
        chapter_label = data.get('chapter')
        question_text = data.get('question_text')

        standards = {
        'L1': 'Problem Solving (L1)',
        'L2': 'Reasoning & Proof (L2)',
        'L3': 'Communication (L3)',
        'L4': 'Connections (L4)',
        'L5': 'Visualization & Representation (L5)'
        }
        standard_label1 = standards[standard_label]       

        #gemini prompt
   
        prompt = """
            Generate a question that is similar in structure and difficulty to the following inputs:\n\n"
            Class: `""" + str(class_label) + """` ,  Subject: `""" + str(subject_label) + """` , Marks:  `""" + str(marks_label) + """` , 
            Standard:  `""" + str(standard_label1) + """`, Chapter:  `""" + str(chapter_label) + """` ,
            Original Question: `""" + str(question_text) + """` , Please ensure the generated question aligns with the topic of the chapter and matches the 
            difficulty level and context of the original question. 
            The output should be a JSON object must be in the following format:
        {
            "question": "New question"
        }
        """

        response = model.generate_content([prompt])
        # print(response.text)
        response_dict = response.text
        # Parse the JSON response
        formatted_str = response_dict.strip("```json\n").strip("```")
        response_data = json.loads(formatted_str)

        # Extract the question
        question = response_data.get("question", "No question found")
        print(question) 




        # Return the AI-generated question as a JSON response
        return jsonify({
            'status': 'success',
            'generated_question': question
        }), 200

    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


# Route to update the question
@app.route('/replaceQuestion/<int:question_id>', methods=['PUT'])
@role_required(['admin', 'teacher'])
def replace_question(question_id):
    try:
        updated_data = request.form
        diagram_file = request.files.get('diagram')
        current_datetime = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')

    
        # Store the diagram file if it exists
        diagram_path = None
        if diagram_file:
            base_name, ext = os.path.splitext(diagram_file.filename)
            new_filename = f"{base_name}_{current_datetime}{ext}"
            diagram_path = os.path.join(DIAGRAMS_FOLDER, new_filename)
            diagram_file.save(diagram_path)
        
        # print(diagram_path)


        # update_question_in_db(question_id, updated_data)
        replace_question_in_db(question_id, {
        'question_text': updated_data['question_text'],
        'diagram': diagram_path  # Pass the diagram path if it exists
        })
        return jsonify({'success': True}), 200
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500



# @app.route('/generatepaper1', methods=['POST'])
# def handle_form_submission():
#     # Retrieve form data
 
#     data = request.get_json()  #
#     selected_class = data.get('class')
#     selected_subject = data.get('subject')
#     selected_assessment = data.get('assessment')
#     selected_chapters = data.get('chapters')
#     sections = data.get('sections')

#     # Log or process the received data
#     print(f"Class: {selected_class}")
#     print(f"Subject: {selected_subject}")
#     print(f"Assessment Type: {selected_assessment}")
#     print(f"Selected Chapters: {selected_chapters}")
#     print(f"Sections: {sections}")

#     # print(data)

#     # Example: Return a success response
#     return jsonify({
#         'message': 'Form submitted successfully'
        
#     })


@app.route('/chapters', methods=['POST'])
@role_required(['admin', 'teacher'])
def fetch_chapters():
    selected_class = request.form['class']
    selected_subject = request.form['subject']
    chapters = get_chapters(selected_class, selected_subject)
    return jsonify(chapters)

# Route for uploading files 
@app.route('/upload', methods=['POST'])
@role_required(['admin', 'teacher'])
def upload_file():
    userid = session.get('user_id')  
    file_id=1
    # Handle file upload
    if 'file' not in request.files:
        return jsonify({"error": "No file  provided"})

    file = request.files['file']

 
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        current_datetime = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')

        # Generate a new unique filename (using the UUID)
        base_name, ext = os.path.splitext(filename)
        new_filename = f"{base_name}_{current_datetime}{ext}"
        
        try:
            file.save(filepath)
        except Exception as e:
            return jsonify({"error": f"Error saving file: {str(e)}"})
        
        # Save file details into the database
        try:
            save_file_to_main_db(new_filename, filename.rsplit('.', 1)[1], filepath,userid)
            file_id = get_file_id_by_filename(new_filename)
            # print(file_id)
            # print(type(file_id))
            file_id = str(file_id)
        except Exception as e:
            return jsonify({"error": f"Error saving file details to database: {str(e)}"})


      
        try:
            # Process file content based on file type
            response_dict=pdf_to_text(filepath)
            # print(response_dict)
            # Remove the markdown block (```json and ```)
            formatted_str = response_dict.strip("```json\n").strip("```")

            # print(type(formatted_str))
            # print(formatted_str)
            # Parse the string into a JSON object
            json_obj = json.loads(formatted_str)
            # Print the resulting JSON object
            # print(json_obj)


            # Check if the JSON data is a list (array) or a dictionary (object)
            if isinstance(json_obj, list):
                # Case: JSON data is a list (array) of sections (first case)
                for section in json_obj:
                    for section,details in section.items():
                        subject1=details['subject']
                        class1=details['class']
                        # class1=None
                        marks1=details['marks']
                        assessment_type=details['assessment']

                        # If any of these values are missing, prompt the user for input
                        if not (subject1 and class1 and assessment_type):
                            delete_file_by_id(file_id)
                            return jsonify({
                                "error": "Missing values",
                                "message": "Please provide class, subject, and assessment type."
                            })
                        
                        for question in details['questions']:
                            insert_question_to_main_db(class1, subject1, assessment_type, marks1, question, file_id)

                assign_chapters(file_id)
                classify_questions(file_id)
                            

            elif isinstance(json_obj, dict):
                # Case: JSON data is a dictionary (object) with named sections (second case)
                for section,details in json_obj.items():
                    subject1=details['subject']
                    class1=details['class']
                    # class1=None
                    marks1=details['marks']
                    assessment_type=details['assessment']

                    # If any of these values are missing, prompt the user for input
                    if not (subject1 and class1 and assessment_type):
                        delete_file_by_id(file_id)
                        return jsonify({
                            "error": "Missing values",
                            "message": "Please provide class, subject, and assessment type."
                        })
                    
                    for question in details['questions']:
                        insert_question_to_main_db(class1, subject1, assessment_type, marks1, question, file_id)

                assign_chapters(file_id)
                classify_questions(file_id)
                        

        except json.JSONDecodeError as e:
            return jsonify({"error": f"Error parsing JSON content: {str(e)}"})



        return jsonify({
            "success": True,
            "message": "File uploaded and processed successfully"
        })


    return jsonify({"error": "Invalid file type or no file provided"})

# Route for uploading files along with data
@app.route('/upload1', methods=['POST'])
@role_required(['admin', 'teacher'])
def upload_file1():
    userid = session.get('user_id')  
    file_id=1
    # Handle file upload
    if 'file' not in request.files:
        return jsonify({"error": "No file  provided"})

    file = request.files['file']

    # Extract form data for class and assessment type
    user_class = request.form.get('class')
    # print(user_class)
    subject = request.form.get('subject')
    # print(subject)
    assessment_type = request.form.get('assessment-type')
    # print(assessment_type)
 
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        current_datetime = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')

        # Generate a new unique filename (using the UUID)
        base_name, ext = os.path.splitext(filename)
        new_filename = f"{base_name}_{current_datetime}{ext}"

        try:
            file.save(filepath)
        except Exception as e:
            return jsonify({"error": f"Error saving file: {str(e)}"})
        
        # Save file details into the database
        try:
            save_file_to_main_db(new_filename, filename.rsplit('.', 1)[1], filepath,userid)
            file_id = get_file_id_by_filename(new_filename)
            # print(file_id)
        except Exception as e:
            return jsonify({"error": f"Error saving file details to database: {str(e)}"})


      
        try:
            # Process file content based on file type
            response_dict=pdf_to_text(filepath)
            # print(response_dict)
            # Remove the markdown block (```json and ```)
            formatted_str = response_dict.strip("```json\n").strip("```")
            # print(type(formatted_str))
            # Parse the string into a JSON object
            json_obj = json.loads(formatted_str)
            # Print the resulting JSON object
            # print(json_obj)

            
            # Check if the JSON data is a list (array) or a dictionary (object)
            if isinstance(json_obj, list):
                # Case: JSON data is a list (array) of sections (first case)
                for section in json_obj:
                    for section,details in section.items():
                        marks1=details['marks']
                        for question in details['questions']:
                            insert_question_to_main_db(user_class, subject, assessment_type, marks1, question, file_id)

                assign_chapters(file_id)
                classify_questions(file_id)

            elif isinstance(json_obj, dict):
                # Case: JSON data is a dictionary (object) with named sections (second case)
                for section,details in json_obj.items():
                    marks1=details['marks']
                    for question in details['questions']:
                        insert_question_to_main_db(user_class, subject, assessment_type, marks1, question, file_id)

                assign_chapters(file_id)
                classify_questions(file_id)

        except json.JSONDecodeError as e:
            return jsonify({"error": f"Error parsing JSON content: {str(e)}"})



        return jsonify({
            "success": True,
            "message": "File uploaded and processed successfully"
        })


    return jsonify({"error": "Invalid file type or no file provided"})

@app.route('/deletequestionpaper/<int:paper_id>', methods=['DELETE'])
@role_required(['admin', 'teacher'])
def delete_question_paper(paper_id):
    try:
        
        delete_questionpaper_from_db(paper_id)
        return jsonify({'message': f'Question paper {paper_id} and its questions have been deleted successfully.'}), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/deletegeneratedpaper/<int:paper_id>', methods=['DELETE'])
@role_required(['admin', 'teacher'])
def delete_generated_question_paper(paper_id):
    try:
        
        delete_generatedpaper_from_db(paper_id)
        filename = f"Question_Paper_{paper_id}.docx"
        filepath = os.path.join(app.config['PAPERS_FOLDER'], filename)
   
        if os.path.exists(filepath):
            os.remove(filepath)
            
        return jsonify({'message': f'Question paper {paper_id} and its questions have been deleted successfully.'}), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500




# Route for /scannedpapers that returns a list of scanned papers in JSON format
@app.route('/scannedpapers', methods=['GET'])
@role_required(['admin', 'teacher'])
def get_scanned_papers():
    role = session.get('role') 
    userid=session.get('user_id')
    if role=='admin':
        files = get_all_files_from_main_db()  # Get the data from the database
        # Convert the result into a list of dictionaries for JSON response
        scanned_papers = []
        for file in files:
            scanned_papers.append({
                'id': file[0],
                'name': file[1],  
                'upload_time': file[4]
            })
        return jsonify(scanned_papers)  # Return as JSON
    else:
        files = get_all_files_based_on_userid_from_main_db(userid)  # Get the data from the database
        # Convert the result into a list of dictionaries for JSON response
        scanned_papers = []
        for file in files:
            scanned_papers.append({
                'id': file[0],
                'name': file[1],  
                'upload_time': file[4]
            })
        return jsonify(scanned_papers)  # Return as JSON





# Route for /scannedpapers that returns a list of scanned papers in JSON format
@app.route('/getQuestionsBasedOnId/<int:id>', methods=['GET'])
@role_required(['admin', 'teacher'])
def get_Questions_basedOn_Id(id):
    questions1 = get_question_based_on_id_from_main_db(id)  # Get the data from the database
    # Convert the result into a list of dictionaries for JSON response
    print(questions1)
    questions=questions1[0]
    Question_data = []
    Question_data.append({
        'id': questions[0],
        'class': questions[1],
        'subject': questions[2],
        'assessment': questions[3],
        'marks': questions[4],
        'question_text': questions[6],
        'diagram':questions[7]
    })
    return jsonify(Question_data)  # Return as JSON




@app.route('/diagrams/<path:filename>')
@role_required(['admin', 'teacher'])
def serve_diagram(filename):
    return send_from_directory('diagrams', filename)



# Route for /getQuestions/<file_id> to fetch questions based on file_id
@app.route('/getQuestions/<file_id>', methods=['GET'])
@role_required(['admin', 'teacher'])
def get_questions(file_id):
    try:
        questions = get_all_questions_based_on_fileid_from_main_db(file_id)  # Get questions based on file_id
        questions_data = []
        
        for question in questions:
            questions_data.append({
                'id': question[0],
                'class': question[1],
                'subject': question[2],
                'assessment': question[3],
                'marks': question[4],
                'question_text': question[6],
                'diagram':question[7]
            })
        
        return jsonify(questions_data)  # Return questions as JSON
    except Exception as e:
        return jsonify({'error': str(e)}), 500  # Handle unexpected error


@app.route('/getAddedQuestions/', methods=['GET'])
@role_required(['admin', 'teacher'])
def get_added_questions():
    user_id = session.get('user_id') 
    role=session.get('role')
    if not user_id:  # Redirect to login if not authenticated
        return redirect(url_for('login'))
    try:
        if role=="admin":
            questions=get_all_questions_with_all_userids_from_main_db()
        else:
            questions = get_all_questions_based_on_userid_from_main_db(user_id)  # Get questions based on user_id
        questions_data = []
        # print("qqq",questions)
        for question in questions:
            questions_data.append({
                'id': question[0],
                'class': question[1],
                'subject': question[2],
                'assessment': question[3],
                'marks': question[4],
                'question_text': question[6],
                'diagram':question[7]
            })
        
        return jsonify(questions_data)  # Return questions as JSON
    except Exception as e:
        print(e)
        return jsonify({'error': str(e)}), 500  # Handle unexpected error



# Route to update the question
@app.route('/editQuestion/<int:question_id>', methods=['PUT'])
@role_required(['admin', 'teacher'])
def edit_question(question_id):
    updated_data = request.form
    diagram_file = request.files.get('diagram')
    current_datetime = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')

   
    # Store the diagram file if it exists
    diagram_path = None
    if diagram_file:
        base_name, ext = os.path.splitext(diagram_file.filename)
        new_filename = f"{base_name}_{current_datetime}{ext}"
        diagram_path = os.path.join(DIAGRAMS_FOLDER, new_filename)
        diagram_file.save(diagram_path)
    
    # print(diagram_path)


    # update_question_in_db(question_id, updated_data)
    update_question_in_db(question_id, {
    'class': updated_data['class'],
    'subject': updated_data['subject'],
    'assessment': updated_data['assessment'],
    'marks': updated_data['marks'],
    'question_text': updated_data['question_text'],
    'diagram': diagram_path  # Pass the diagram path if it exists
})
    return jsonify({'success': True}), 200


@app.route('/deletequestionbasedonId/<int:id>', methods=['DELETE'])
@role_required(['admin', 'teacher'])
def delete_question_id(id):
    try:
        
        delete_question_from_db(id)
        return jsonify({'message': f'Question  {id}  deleted successfully.'}), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500


# Function to save question to Main_db
@app.route('/AddQuestions', methods=['POST'])
@role_required(['admin', 'teacher'])
def update_questions():
    userid = session.get('user_id')  
    current_datetime = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
    
    try:
        class_name = request.form['class']
        subject = request.form['subject']
        assessment = request.form['assessment']
        marks = request.form['marks']
        question_text = request.form['question']
        diagram_file = request.files.get('diagram')  # Get the diagram file if uploaded
    except KeyError as e:
        return jsonify({"error": f"Missing key: {str(e)}"}), 400
    
    if not class_name or not subject or not assessment or not marks or not question_text:
        return jsonify({"error": "All fields are required."}), 400

    # Check if the question already exists
    existing_question = check_question_exists_in_main_db(class_name, subject, assessment, marks, question_text)
    
    if existing_question:
        return jsonify({"error": "This question already exists in the database."}), 400
    else:
        # Handle the diagram file upload
        diagram_path = None
        if diagram_file:
            base_name, ext = os.path.splitext(diagram_file.filename)
            new_filename = f"{base_name}_{current_datetime}{ext}"
            diagram_path = os.path.join(DIAGRAMS_FOLDER, new_filename)
            diagram_file.save(diagram_path)

        # Insert the new question into the Main_db
        insert_single_question_to_main_db(class_name, subject, assessment, marks, question_text,userid,diagram_path)
        return jsonify({"message": "Question added successfully."}), 200

def pdf_to_text(file_path):
    sample_file = genai.upload_file(path=file_path, display_name="pdf1")

    print(f"Uploaded file '{sample_file.display_name}' as: {sample_file.uri}")

    
    
    response = model.generate_content([
        sample_file,
        "I need you to extract and list the questions from each section of the provided PDF document. For each section, you should output a JSON object like this ***mandatory format***: for example like this: { \"section-1\": {\"questions\": [\"question1\", \"question2\"], \"subject\": \"Mathematics\", \"class\": \"10\", \"marks\": \"2\",\"assessment\": \"summative\"}}. For each section, the subject, class in digital number not roman number , assessment and marks for each question not total marks should be provided, if any value not available in pdf return None. If the questions are multiple-choice, include the options within the question like this: A) option1 B) option2 C) option3 D) option4. The output must strictly follow this format for each section, and it should be in JSON format.For part-B include it as section-n in the output and dont include part-A in output , include sections of Part-A" ,

        
    ])

    # Print the generated content
    # print(response.text)

    return response.text

def assign_chapters_to_questions(questions_list, chapter_list):
   
    prompt = """
    I have a list of questions and a list of chapters. I need you to assign the most relevant chapter name from the provided chapter list to each question. 
    Here are the chapters: `""" + str(chapter_list) + """` , Here are the questions: `""" + str(questions_list) + """`, 
    For each question, analyze its content and context to determine the best matching chapter from the chapter list. The assigned chapter must be one of the provided chapters. Do not create or suggest new chapters.
    The output should be a JSON object in the following format:
      [
        {
            "id": "1",
            "question": "question1",
            "chapter": "Chapter X - Chapter Name"
        },
        {
            "id": "2",
            "question": "question2",
            "chapter": "Chapter Y - Chapter Name"
        }
    ]  

    """

    response = model.generate_content([prompt])
    return response.text


def assign_chapters(file_id):
    chapter_list = [
        "Chapter 1 - Real Numbers",
        "Chapter 2 - Sets",
        "Chapter 3 - Polynomials",
        "Chapter 4 - Pair of Linear Equations in Two Variables",
        "Chapter 5 - Quadratic Equations",
        "Chapter 6 - Progressions",
        "Chapter 7 - Coordinate Geometry",
        "Chapter 8 - Similar Triangles",
        "Chapter 9 - Tangents and Secants to a Circle",
        "Chapter 10 - Mensuration",
        "Chapter 11 - Trigonometry",
        "Chapter 12 - Applications of Trigonometry",
        "Chapter 13 - Probability",
        "Chapter 14 - Statistics"
    ]
    
    questions_list = get_only_questions_based_on_fileid_from_main_db(file_id)

    response_dict=assign_chapters_to_questions(questions_list,chapter_list)

    # print(response_dict)
    # Remove the markdown block (```json and ```)
    formatted_str = response_dict.strip("```json\n").strip("```")
    # print(type(formatted_str))
    # Parse the string into a JSON object
    json_obj = json.loads(formatted_str)

    # Print the resulting JSON object
    # print(json_obj)

    # Extract `id`, `question`, and `chapter`
    for item in json_obj:
        question_id = item["id"]
        question_text = item["question"]
        chapter_name = item["chapter"]

        # Print or process the extracted data
        # print(f"ID: {question_id}")
        # print(f"Question: {question_text}")
        # print(f"Chapter: {chapter_name}")
        # print("-" * 40)

        insert_chapters_to_main_db(question_id,question_text,chapter_name)

def classify_questions_to_academic_standards(questions_list, academic_standards):
    prompt = f"""
    I have a list of questions and a list of academic standards. I need you to classify each question into the most relevant academic standard based on its content. 
    Here are the academic standards: {academic_standards}
    Here are the questions: {questions_list}
    For each question, analyze its content and context to determine the best matching academic standard from the list. Give only L1,L2,L3,L4,L5, dont give full name. 
    The output should be a JSON object must be in the following format:
    [
        {{
            "id": "1",
            "question": "question1",
            "academic_standard": "L1"
        }},
        {{
            "id": "2",
            "question": "question2",
            "academic_standard": "L2"
        }}
    ]
    """
    
    response = model.generate_content([prompt])
    # print("Raw response:", response.text)  # Debugging step
    return response.text

def classify_questions(file_id):
    # List of academic standards
    academic_standards = [
        "Problem Solving (L1)",
        "Reasoning & Proof (L2)",
        "Communication (L3)",
        "Connections (L4)",
        "Visualization & Representation (L5)"
    ]

    # Retrieve the list of questions from the database
    questions_list = get_only_questions_based_on_fileid_from_main_db(file_id)

    if not questions_list:
        print("Error: No questions retrieved from the database.")
    else:
        # Call the function to classify questions into academic standards
        response_str = classify_questions_to_academic_standards(questions_list, academic_standards)

        if response_str:
            try:
                # print("Response string:", response_str)  # Debugging step

                # Remove the markdown block (```json and ```)
                formatted_str = response_str.strip("```json\n").strip("```")
                # print(type(formatted_str))
                # Parse the string into a JSON object
                json_obj = json.loads(formatted_str)

                # Extract id, question, and academic_standard
                for item in json_obj:
                    question_id = item.get("id")
                    academic_standard = item.get("academic_standard")

                    if question_id and academic_standard:
                        # Print or process the extracted data
                        # print(f"ID: {question_id}")

                        # print(f"Academic Standard: {academic_standard}")
                        # print("-" * 40)

                        # Insert the results into the database (add a column for academic standards)
                        insert_academic_standard_to_main_db(question_id, academic_standard)
                    else:
                        print(f"Missing data in item: {item}")

            except json.JSONDecodeError as e:
                print(f"Error parsing JSON: {e}")
        else:
            print("Error: No response from the Generative AI API.")


def save_question_paper_to_word(sections,filename,folder):
    """
    Save the question paper to a Word document.
    """
    document = Document()
    document.add_heading("Question Paper", level=1)

    for section_num, section in enumerate(sections):
        # print(section_num)
        # print(section)
        document.add_heading(f"Section {section_num +1}: {section['marks']} Marks", level=2)
        for question_num, question in enumerate(section['questions']):
            # print(question_num+1,question)
            question_text = question[1]
            chapter = question[3]
            standard = question[4]
            document.add_paragraph(f"{question_num +1}. {question_text} (Chapter: {chapter}, Standard: {standard})")

    # Ensure the upload folder exists
    if not os.path.exists(folder):
        os.makedirs(folder)


    filepath = os.path.join(folder, filename)
    document.save(filepath)
    # print(filepath)
    return  filepath


@app.route('/downloadpaper/<paper_id>', methods=['GET'])
@role_required(['admin', 'teacher'])
def download_question_paper(paper_id):
    try:
       
        questions = get_all_questions_based_on_paperid_from_main_db(paper_id)
        # print(questions)


        # Organize questions into sections based on `sectionid`
        sections_dict = {}
        for question in questions:
             # Extract data from the question
            question_id, question_text, chapter, section_id, marks, diagram, standard, paper_id = question
            
            if section_id not in sections_dict:
                sections_dict[section_id] = {
                    "marks": marks,  # Use the marks from the first question in this section
                    "questions": []
                }
            
            # Append question details
            sections_dict[section_id]['questions'].append([
                question_id,
                question_text,
                diagram,
                chapter,
                standard
            ])
        
        # Convert dictionary to a list of sections
        sections = [
            {"marks": section_data["marks"], "questions": section_data["questions"]}
            for section_id, section_data in sections_dict.items()
        ]
        # print(sections)
        # Generate and save the Word document
        filename = f"Question_Paper_{paper_id}.docx"
        file_path = save_question_paper_to_word(sections, filename, app.config['PAPERS_FOLDER'])

        return send_file(file_path, as_attachment=True, download_name=filename)
    

    except Exception as e:
        return jsonify({'error': str(e)}), 500  # Handle unexpected errors
        

@app.route('/generatepaper1', methods=['POST'])
@role_required(['admin', 'teacher'])
def handle_form_submission():
    # Retrieve form data
 
    data = request.get_json()  #
    selected_class = data.get('class')
    selected_subject = data.get('subject')
    selected_assessment = data.get('assessment')
    selected_chapters = data.get('chapters')
    selected_sections = data.get('sections')

    # Log or process the received data
    # print(f"Class: {selected_class}")
    # print(f"Subject: {selected_subject}")
    # print(f"Assessment Type: {selected_assessment}")
    # print(f"Selected Chapters: {selected_chapters}")
    # print(f"Sections: {selected_sections}")

    num_sections = len(selected_sections)
    sections = []

    for i in range(num_sections):
        # print(f"\nSection {i+1}:")
        marks = selected_sections[i]['marks']
        num_questions = selected_sections[i]['numQuestions']
        # Access standards
        standards_list = []
        standard_distribution = {}
        standards = selected_sections[i]['standards']
        for standard in standards:
            standard_name = standard['standard']
            standards_list.append(standard_name)
            num_ques = standard['num_ques']
            standard_distribution[standard_name] = int(num_ques)

        # print("std",standard_distribution)
        # print("stl",standards_list)
        # Fetch questions for this section
        questions = fetch_questions(selected_class, selected_subject, selected_assessment, selected_chapters, standards_list)

        if len(questions) < int(num_questions):
            print(f"Only {len(questions)} questions available for this section. Adjusting to available questions.")
            num_questions = len(questions)

        # If the fetched questions exceed the required number, randomly sample them
        # if len(questions) > int(num_questions):
        #     questions = random.sample(questions, int(num_questions))

        selected_questions = distribute_questions_by_standard(questions, standard_distribution, num_questions)
        sections.append({
            'marks': marks,
            'questions': selected_questions
        })
        # print("sections",sections)

    # Step 2: send question paper to database
    created_time = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
    userid = session.get('user_id')  
    paper_id = insert_generatedpaper_to_main_db(selected_class, selected_subject, selected_assessment, num_sections, created_time,userid)

    
    for section_num, section in enumerate(sections, start=1):
        #print(f"\nSection {section_num}: {section['marks']} Marks")
        for question_num, question in enumerate(section['questions'], start=1):
            question_text = question[1]
            chapter = question[3]
            standard = question[4]
            insert_generatedquestion_to_main_db(question_text,chapter, section_num, section['marks'], standard , paper_id )


    # Example: Return a success response
    return jsonify({
        'message': 'Form submitted successfully'
        
    })

def process_attendance_data(file_path):
    try:
        df = pd.read_csv(file_path)
        total_row_index = df[df[df.columns[0]].str.contains("Total", na=False)].index[0]
        df = df.iloc[1:total_row_index]

        df.columns = ['Class', 'Students'] + df.columns[2:].tolist()
        df['Students'] = pd.to_numeric(df['Students'], errors='coerce')

        date_columns = df.columns[2:]
        for col in date_columns:
            df[col + '_Holiday_Flag'] = df[col].isnull().astype(int)
            df[col] = df[col].fillna(0)

        date_columns = [col for col in date_columns if df[col].notnull().any()]
        new_column_names = ['Day' + str(i + 1) for i in range(len(date_columns))]
        df.rename(columns=dict(zip(date_columns, new_column_names)), inplace=True)

        percent_cols = []
        for day in new_column_names:
            percent_col = day + '_percent'
            df[percent_col] = (df[day] / df['Students']) * 100
            percent_cols.append(percent_col)

        scaler = MinMaxScaler()
        df[percent_cols] = scaler.fit_transform(df[percent_cols])
        

        df['Average_Attendance'] = df[percent_cols].mean(axis=1)
        df['Attendance_Deviation'] = df[percent_cols].std(axis=1)
        df['Consistent_Attendance_Flag'] = df['Attendance_Deviation'] < 0.1
        df['Average_Attendance']=df['Average_Attendance']*100

        df_sorted = df.sort_values(by='Average_Attendance', ascending=False)
        return df_sorted, percent_cols
    except Exception as e:
        raise ValueError(f"Error processing file: {str(e)}")



@app.route('/attendence', methods=['GET', 'POST'])
@role_required(['admin', 'teacher'])
def attendence():
    role = session.get('role')  
    if not role:  # Redirect to login if not authenticated
        return redirect(url_for('login'))
    if request.method == 'POST':
        num_months = int(request.form.get('months', 0))
        files = []
        for i in range(1, num_months + 1):
            file = request.files.get(f'file{i}')
            if file and allowed_file(file.filename):
                file_path = os.path.join(app.config['EXCELUPLOAD_FOLDER'], file.filename)
                file.save(file_path)
                files.append(file_path)

        if not files:
            return render_template('attendence.html', error="No files uploaded or invalid file types.",role=role)

        try:
            dfs = []
            for file_path in files:
                df_sorted, percent_cols = process_attendance_data(file_path)
                dfs.append(df_sorted)

            combined_df = pd.concat(dfs, ignore_index=True)
            combined_df = combined_df.groupby('Class').mean().reset_index()
            combined_df.sort_values(by='Average_Attendance', ascending=False, inplace=True)

            output_file = os.path.join(app.config['EXCELUPLOAD_FOLDER'], 'attendance_analysis.xlsx')
            combined_df.to_excel(output_file, index=False)

            result = combined_df[['Class', 'Average_Attendance']].to_dict('records')

            return render_template('attendence.html', uploaded=True, result=result,role=role)
        except ValueError as e:
            return render_template('attendence.html', error=str(e),role=role)

    return render_template('attendence.html',role=role)


def analyze_class_data(filepath, class_name, passing_marks):
    try:
        df = pd.read_excel(filepath)
    except FileNotFoundError:
        return None, "File not found."

    student_columns = df.columns[1:-1]
    if not student_columns.size:
        return None, "No valid student columns found in the Excel file."

    df["Total"] = df[student_columns].sum(axis=1)
    df["Average"] = df[student_columns].mean(axis=1)
    student_averages = df[student_columns].mean().sort_values(ascending=False)

    if student_averages.empty or student_averages.isnull().all():
        return None, "No valid student data found."

    topper = student_averages.idxmax()
    sorted_students = student_averages.reset_index()
    sorted_students.columns = ["Student", "Average Marks"]
    sorted_students["Topper"] = sorted_students["Student"].apply(lambda x: "Yes" if x == topper else "No")

    class_average = sorted_students["Average Marks"].mean()
    passing_students = sorted_students[sorted_students["Average Marks"] >= passing_marks]
    failing_students = sorted_students[sorted_students["Average Marks"] < passing_marks]
    total_students = len(sorted_students)
    pass_percentage = (len(passing_students) / total_students * 100) if total_students > 0 else 0
    fail_percentage = (len(failing_students) / total_students * 100) if total_students > 0 else 0

    chart_img = io.BytesIO()
    labels = 'Pass', 'Fail'
    sizes = [pass_percentage, fail_percentage]
    explode = (0.1, 0)
    plt.pie(sizes, explode=explode, labels=labels, autopct='%1.1f%%', shadow=True, startangle=90)
    plt.axis('equal')
    plt.title(f"Pass/Fail Percentage for {class_name} (Passing Marks: {passing_marks})")
    plt.savefig(chart_img, format='png')
    plt.close()
    chart_img.seek(0)
    pie_chart_url = base64.b64encode(chart_img.getvalue()).decode()

    return {
        "sorted_students": sorted_students,
        "class_average": class_average,
        "pass_percentage": pass_percentage,
        "fail_percentage": fail_percentage,
        "pie_chart_url": pie_chart_url
    }, None


@app.route('/classperformance', methods=['GET', 'POST'])
@role_required(['admin', 'teacher'])
def class_performance():
    role = session.get('role')  
    if not role:  # Redirect to login if not authenticated
        return redirect(url_for('login'))
    if request.method == 'POST':
        class_name = request.form.get('class_name', '')
        try:
            passing_marks = int(request.form.get('passing_marks', '0'))
        except ValueError:
            return render_template('ClassPerformance.html', error="Invalid passing marks.",role=role)

        file = request.files.get('file')
        if file and allowed_file(file.filename):
            filepath = os.path.join(app.config['EXCELUPLOAD_FOLDER'], file.filename)
            file.save(filepath)

            analysis_result, error = analyze_class_data(filepath, class_name, passing_marks)
            if error:
                return render_template('ClassPerformance.html', error=error,role=role)

            return render_template(
                'ClassPerformance.html',
                analysis=True,
                sorted_students=analysis_result["sorted_students"].to_dict(orient="records"),
                class_average=analysis_result["class_average"],
                pass_percentage=analysis_result["pass_percentage"],
                fail_percentage=analysis_result["fail_percentage"],
                pie_chart_url=analysis_result["pie_chart_url"],
                role=role
            )

        return render_template('ClassPerformance.html', error="Invalid file type or no file uploaded.",role=role)

    return render_template('ClassPerformance.html',role=role)




def analyze_feedback(filepath, passing_marks, class_name):
    df = pd.read_excel(filepath)

    if 'Teachers' not in df.columns or 'Subject' not in df.columns:
        raise ValueError("The file must contain 'Teachers' and 'Subject' columns.")

    teacher_feedback = []

    for _, row in df.iterrows():
        teacher = row['Teachers']
        subject = row['Subject']
        student_scores = row[1:-1]  # Ignore Subject and Teachers columns
        total_students = len(student_scores)
        passing_students = student_scores[student_scores >= passing_marks].count()
        feedback_percentage = (passing_students / total_students * 100) if total_students > 0 else 0

        teacher_feedback.append({
            'teacher_name': teacher,
            'subject_name': subject,
            'feedback_percentage': feedback_percentage,
            'class_name': class_name
        })

        insert_teacherfeedback(teacher, subject, feedback_percentage, class_name)
        

    # Visualization
    chart_img = io.BytesIO()
    teachers_names = [feedback['teacher_name'] for feedback in teacher_feedback]
    feedback_percentages = [feedback['feedback_percentage'] for feedback in teacher_feedback]

    plt.barh(teachers_names, feedback_percentages, color='skyblue')
    plt.xlabel('Feedback Percentage')
    plt.ylabel('Teachers')
    plt.title(f'Teacher Feedback Analysis')
    plt.tight_layout()
    plt.savefig(chart_img, format='png')
    plt.close()
    chart_img.seek(0)
    bar_chart_url = base64.b64encode(chart_img.getvalue()).decode()

    return teacher_feedback, bar_chart_url



@app.route('/teacherperformance', methods=['GET', 'POST'])
@role_required(['admin'])
def teacherperformance():
    role = session.get('role')  
    if not role:  # Redirect to login if not authenticated
        return redirect(url_for('login'))
    if request.method == 'POST':
        try:
            passing_marks = int(request.form.get('passing_marks', '0'))
            class_name = request.form.get('class_name', '').strip()
            if not class_name:
                raise ValueError("Class name is required.")
        except ValueError as e:
            return render_template('TeacherPerformance.html', error=str(e),role=role)

        file = request.files.get('file')
        if file and allowed_file(file.filename):
            filepath = os.path.join(app.config['EXCELUPLOAD_FOLDER'], file.filename)
            file.save(filepath)

            try:
                teacher_feedback, bar_chart_url = analyze_feedback(filepath, passing_marks, class_name)
                return render_template(
                    'TeacherPerformance.html',
                    analysis=True,
                    teacher_feedback=teacher_feedback,
                    bar_chart_url=bar_chart_url
                    ,role=role
                )
            except Exception as e:
                return render_template('TeacherPerformance.html', error=str(e),role=role)

        return render_template('TeacherPerformance.html', error="Invalid file type or no file uploaded.",role=role)

    return render_template('TeacherPerformance.html',role=role)

@app.route('/teacherfeedback', methods=['GET'])
@role_required(['admin'])
def teacher_feedback():
    role = session.get('role')  
    if not role:  # Redirect to login if not authenticated
        return redirect(url_for('login'))
    data = get_teacherfeedback_from_main_db()
    return {'teacher_feedback': data,'role':role}






if __name__ == '__main__':
    app.run(debug=True)